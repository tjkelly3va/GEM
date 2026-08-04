#!/usr/bin/env python3
"""
gem_policy_to_docx.py — export one GEM policy to a Microsoft Word (.docx) readout.

Given a policy identifier such as "NCD 30.3", this resolves the matching
individual in the GEM graph and writes a Word document containing:

  * the policy individual (gem:NCDpolicy / any gem:CMSpolicy) as a curated
    readout — identity, dates, publication, benefit categories, clinical
    concepts, provider credentials, healthcare settings, and outbound
    references.  The internal audit narrative (gem:workflowDescription) is
    intentionally omitted.
  * its policy-level rules            (gem:hasPolicyRule      -> gem:PolicyRule)
  * any linked Policy Groups          (gem:hasPolicyGroup     -> gem:PolicyGroup)
  * any linked Policy Coding Rules    (gem:hasPolicyCodingRule-> gem:PolicyCodingRule)

Policy Groups and Policy Coding Rules are rendered with their code anchors
(ICD-10 / HCPCS / CPT) and with any rules nested on them via gem:hasPolicyRule.

Design decisions (confirmed by Tom, S235):
  * EXACT identifier match only — "NCD 30.3" does NOT pull in 30.3.1/.2/.3.
  * Rules nested under a Group / Coding Rule ARE included, under their scope.
  * The policy individual is a curated readout; gem:workflowDescription omitted.

Resolution is by gem:identifier (whitespace-collapsed, case-insensitive), which
is more robust than reconstructing a URI from the label.

Usage (from the repo root):
    python .claude/skills/gem-policy-docx/scripts/gem_policy_to_docx.py \
        --identifier "NCD 30.3" \
        --files-dir .claude/skills/policy-extraction \
        --output export/NCD_30.3.docx

If --output is omitted a filename is derived from the identifier and written to
the current directory. If --files-dir is omitted it defaults to the current
directory, so invoking from inside the canonical directory needs no flag.
"""

import argparse
import os
import re
import sys

import rdflib
from rdflib import RDF, URIRef, Literal
from rdflib.namespace import Namespace, XSD

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# GEM / GEMI are RUNTIME-DETECTED from the loaded graph's own @prefix
# declarations (see detect_gem_namespace + bind_gem_namespace), so a namespace
# version bump needs no edit here. The literals below are only the fallback used
# when no declaration is found; keep them at the most recent known namespace so a
# detection miss still degrades to a working default. This mirrors gem_audit.py's
# Option B (S255).
#
# Re-hardcoding these is the S260 defect and it fails SILENTLY-shaped: under a
# stale namespace every gem: query returns zero rows, which is indistinguishable
# from "no such policy", so the script exits telling you to check an identifier
# that was correct all along — for every policy, not just the one you asked for.
GEM = Namespace("http://www.cms.hhs.gov/ontology/2026/08/GEM/")
GEMI = Namespace("http://www.cms.hhs.gov/ontology/2026/08/GEM/instances/")
DC = Namespace("http://purl.org/dc/elements/1.1/")

# External code vocabularies -> (system label, URI stem)
CODE_SYSTEMS = [
    ("ICD-10-CM", "http://purl.bioontology.org/ontology/ICD10CM/"),
    ("HCPCS", "http://purl.bioontology.org/ontology/HCPCS/"),
    ("CPT", "https://www.ama-assn.org/cpt#"),
    ("CPT", "http://purl.bioontology.org/ontology/CPT/"),
]

# Canonical instance files. Others are tolerated but these carry the policy web.
DEFAULT_FILES = [
    "GEM_policy_instances.ttl",
    "GEM_code_group_instances.ttl",
    "cpt.ttl",
]

# These two are the ONLY module-level constants derived from GEM, so they are the
# only ones runtime detection has to rebuild. They are built inside a function
# rather than at import time on purpose: a GEM.-qualified URIRef captured at
# import freezes the namespace, and rebinding the GEM global afterwards would
# leave these still pointing at the old one — a rebind that looks like it worked
# and silently doesn't. gem_audit.py carries the scar of exactly this
# (_LLM_ANNOTATED_VOCAB_LOCALNAMES holds local names, not URIRefs, for the same
# reason). bind_gem_namespace() calls this again after detection.
POLICY_IDENTITY_PREDS: list = []
SUPPRESS_PREDS: set = set()


def _rebuild_gem_derived() -> None:
    """Rebuild every GEM-derived module constant from the current GEM binding."""
    global POLICY_IDENTITY_PREDS, SUPPRESS_PREDS
    # Predicates handled explicitly in the curated policy readout, so the generic
    # "other references" pass skips them.
    POLICY_IDENTITY_PREDS = [
        ("Identifier", GEM.identifier),
        ("Version", GEM.policyVersion),
        ("Effective date", GEM.policyEffectiveDate),
        ("Implementation date", GEM.policyImplementationDate),
        ("Publication number", GEM.publicationNumber),
        ("Manual section", GEM.manualSectionNumber),
        ("Page count", GEM.policyPageCount),
        ("In effect", GEM.isInEffect),
        ("Next planned step", GEM.nextPlannedStep),
    ]
    # Predicates never shown as raw rows (structural / provenance / elsewhere).
    SUPPRESS_PREDS = {
        RDF.type, GEM.memberOfOntology, GEM.prefLabel, GEM.description,
        GEM.workflowDescription, GEM.identifier, GEM.hasPolicyRule,
        GEM.hasPolicyGroup, GEM.hasPolicyCodingRule, DC.source,
    }
    SUPPRESS_PREDS |= {p for _, p in POLICY_IDENTITY_PREDS}


def detect_gem_namespace(g: rdflib.Graph) -> tuple[str, str]:
    """Read the GEM / GEMI base IRIs from the parsed graph's own prefix bindings.

    rdflib preserves each `@prefix` declaration it parsed, so the loaded corpus
    is self-describing and no file re-read is needed. `gem:` is read directly;
    `gemi:` is read directly when present, else derived as gem + "instances/".
    Falls back to the module-level literals when neither is declared.
    """
    bound = {p: str(n) for p, n in g.namespaces()}
    gem_uri = bound.get("gem")
    if gem_uri is None:
        return str(GEM), str(GEMI)
    return gem_uri, bound.get("gemi", gem_uri + "instances/")


def bind_gem_namespace(g: rdflib.Graph) -> str:
    """Detect and bind the namespace for this run. Returns the GEM base IRI."""
    global GEM, GEMI
    gem_uri, gemi_uri = detect_gem_namespace(g)
    GEM, GEMI = Namespace(gem_uri), Namespace(gemi_uri)
    _rebuild_gem_derived()
    return gem_uri


_rebuild_gem_derived()

# Friendly names for common predicates; fall back to de-camelCased local name.
PRED_LABELS = {
    "refersToBenefitCategory": "Benefit category",
    "refersToClinicalConcept": "Clinical concept",
    "requiresProviderCredential": "Provider credential",
    "requiresHealthcareSetting": "Healthcare setting",
    "refersToHealthcareSetting": "Healthcare setting",
    "referencesPolicy": "References policy",
    "revisesPolicy": "Revises policy",
    "referencesChangeRequest": "References change request",
    "transmitsChangeRequest": "Transmits change request",
    "appliesWhenCode": "Applies when code",
    "appliesWhenProcedure": "Applies when procedure (HCPCS)",
    "appliesWhenDiagnosis": "Applies when diagnosis (ICD-10)",
    "coversCondition": "Covers condition (ICD-10)",
    "coversProcedure": "Covers procedure (HCPCS)",
    "exclusivelyCoversCondition": "Exclusively covers condition (ICD-10)",
    "exclusivelyCoversPrimaryCondition": "Exclusively covers primary condition (ICD-10)",
    "exclusivelyCoversSecondaryCondition": "Exclusively covers secondary condition (ICD-10)",
    "exclusivelyCoversProcedure": "Exclusively covers procedure (HCPCS)",
    "excludesCondition": "Excludes condition (ICD-10)",
    "excludesProcedure": "Excludes procedure (HCPCS)",
    "refersToICDdiagnosis": "Refers to ICD-10 diagnosis",
    "refersToHCPCSprocedure": "Refers to HCPCS procedure",
    "memberOfCodeGroup": "Member of code group",
    "hasCodeGroup": "Code group",
}


def local(uri):
    s = str(uri)
    for sep in ("#", "/"):
        if sep in s:
            s = s.rsplit(sep, 1)[-1]
    return s


def pred_label(p):
    ln = local(p)
    if ln in PRED_LABELS:
        return PRED_LABELS[ln]
    # de-camelCase: refersToX -> "Refers to x"
    words = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", ln)
    return words[:1].upper() + words[1:]


def code_render(uri):
    """Render an external code URI as 'CODE (System)', else None."""
    s = str(uri)
    for label, stem in CODE_SYSTEMS:
        if s.startswith(stem):
            return f"{s[len(stem):]} ({label})"
    return None


def obj_render(g, o):
    """Human-readable rendering of an object node."""
    if isinstance(o, Literal):
        return str(o)
    code = code_render(o)
    if code:
        return code
    # instance in the graph: prefer identifier, then prefLabel, then local name
    ident = g.value(o, GEM.identifier)
    if ident:
        label = g.value(o, GEM.prefLabel)
        return f"{ident}" + (f" — {label}" if label and str(label) != str(ident) else "")
    label = g.value(o, GEM.prefLabel)
    if label:
        return str(label)
    return local(o)


def sort_key(text):
    # natural-ish sort: split digit runs so 2 < 10
    return [int(t) if t.isdigit() else t.lower()
            for t in re.split(r"(\d+)", text)]


# --------------------------------------------------------------------------- #
# Turtle serialization (for the triples appendix)
# --------------------------------------------------------------------------- #

# namespace URI -> prefix.  gemi is a sub-namespace of gem, so match longest first.
PREFIXES = {
    "http://www.w3.org/1999/02/22-rdf-syntax-ns#": "rdf",
    "http://www.w3.org/2000/01/rdf-schema#": "rdfs",
    "http://www.w3.org/2002/07/owl#": "owl",
    "http://www.w3.org/2004/02/skos/core#": "skos",
    "http://www.w3.org/ns/shacl#": "sh",
    "http://purl.org/dc/elements/1.1/": "dc",
    "http://www.w3.org/2001/XMLSchema#": "xsd",
    "http://purl.bioontology.org/ontology/HCPCS/": "hcpcs",
    "http://purl.bioontology.org/ontology/ICD10CM/": "icd10",
    "https://www.ama-assn.org/cpt#": "cpt",
    "http://www.cms.hhs.gov/ontology/2026/07/GEM/instances/": "gemi",
    "http://www.cms.hhs.gov/ontology/2026/07/GEM/": "gem",
    "http://www.cms.hhs.gov/ontology/2026/07/RBCS/": "rbcs",
}
_NS_BY_LEN = sorted(PREFIXES, key=len, reverse=True)


def ttl_qname(uri):
    s = str(uri)
    for ns in _NS_BY_LEN:
        if s.startswith(ns):
            return f"{PREFIXES[ns]}:{s[len(ns):]}"
    return f"<{s}>"


def ttl_literal(lit):
    val = str(lit)
    if lit.datatype == XSD.integer:
        return val
    esc = (val.replace("\\", "\\\\").replace('"', '\\"')
              .replace("\n", "\\n").replace("\r", "\\r").replace("\t", "\\t"))
    if lit.language:
        return f'"{esc}"@{lit.language}'
    if lit.datatype is not None:
        return f'"{esc}"^^{ttl_qname(lit.datatype)}'
    return f'"{esc}"'


def ttl_term(node):
    if isinstance(node, Literal):
        return ttl_literal(node)
    if isinstance(node, URIRef):
        return ttl_qname(node)
    return f"_:{node}"   # blank node (none expected in this graph)


def render_subject_ttl(g, subj, depth, lines, emitted):
    """Append indented Turtle statements for `subj` at nesting `depth`.
    Each subject is emitted at most once."""
    if subj in emitted:
        return
    emitted.add(subj)
    ind = " " * (4 * depth)
    pind = " " * (4 * depth + 4)

    def order(po):
        p, _o = po
        if p == RDF.type:
            return (0, "")
        if p == DC.source:
            return (2, "")
        return (1, ttl_qname(p) + " " + ttl_term(_o))

    pos = sorted((po for po in g.predicate_objects(subj)
                  if po[0] != GEM.workflowDescription), key=order)
    stmts = [("a" if p == RDF.type else ttl_qname(p), ttl_term(o)) for p, o in pos]
    if not stmts:
        return
    subjq = ttl_qname(subj)
    n = len(stmts)
    for i, (pred, obj) in enumerate(stmts):
        end = " ." if i == n - 1 else " ;"
        if i == 0:
            lines.append(f"{ind}{subjq} {pred} {obj}{end}")
        else:
            lines.append(f"{pind}{pred} {obj}{end}")
    lines.append("")   # blank line between subject blocks


# --------------------------------------------------------------------------- #
# Word helpers
# --------------------------------------------------------------------------- #

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)   # deep slate blue
MUTED = RGBColor(0x55, 0x5B, 0x66)


def set_letter(doc):
    for sec in doc.sections:
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11)
        sec.left_margin = sec.right_margin = Inches(1)
        sec.top_margin = sec.bottom_margin = Inches(1)


def style_base(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for lvl, size in ((1, 17), (2, 13), (3, 11.5)):
        st = doc.styles[f"Heading {lvl}"]
        st.font.size = Pt(size)
        st.font.color.rgb = ACCENT
        st.font.name = "Calibri"


def add_kv_table(doc, rows):
    """Two-column key/value table."""
    if not rows:
        return
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Light Grid Accent 1"
    for k, v in rows:
        c = t.add_row().cells
        c[0].width = Inches(2.2)
        c[1].width = Inches(4.3)
        kp = c[0].paragraphs[0]
        run = kp.add_run(k)
        run.bold = True
        c[1].paragraphs[0].add_run(v)
    doc.add_paragraph()


def add_bullets(doc, items, label=None):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_muted(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = MUTED


_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def emit_code_line(doc, text):
    """One monospace paragraph, leading spaces preserved, tight spacing."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    if text:
        tel = run._element.find(qn("w:t"))
        if tel is not None:
            tel.set(_XML_SPACE, "preserve")


# --------------------------------------------------------------------------- #
# Content builders
# --------------------------------------------------------------------------- #

def grouped_objects(g, subj, skip=SUPPRESS_PREDS):
    """Return {predicate_label: [rendered objects sorted]} for a subject,
    skipping suppressed predicates."""
    buckets = {}
    for p, o in g.predicate_objects(subj):
        if p in skip:
            continue
        buckets.setdefault(p, []).append(obj_render(g, o))
    out = {}
    for p, vals in buckets.items():
        out[pred_label(p)] = sorted(set(vals), key=sort_key)
    return out


def render_rules(doc, g, rules, heading, level):
    if not rules:
        return
    doc.add_heading(heading, level=level)
    for r in sorted(rules, key=lambda x: sort_key(local(x))):
        label = g.value(r, GEM.prefLabel) or local(r)
        p = doc.add_paragraph()
        p.add_run(str(label)).bold = True
        rtypes = sorted(local(t) for t in g.objects(r, GEM.ruleType))
        if rtypes:
            rt = ", ".join(t.replace("ruleType_", "") for t in rtypes)
            rp = doc.add_paragraph()
            run = rp.add_run(f"Rule type: {rt}")
            run.italic = True
            run.font.color.rgb = MUTED
        desc = g.value(r, GEM.ruleDescription)
        if desc:
            doc.add_paragraph(str(desc))


def render_scope(doc, g, scope, kind, level):
    """Render a PolicyGroup or PolicyCodingRule with its anchors + nested rules."""
    label = g.value(scope, GEM.prefLabel) or local(scope)
    doc.add_heading(f"{kind}: {label}", level=level)
    desc = g.value(scope, GEM.description)
    if desc:
        doc.add_paragraph(str(desc))
    # code anchors and any other object properties
    buckets = grouped_objects(g, scope)
    for plabel in sorted(buckets):
        add_bullets(doc, buckets[plabel], label=plabel)
    # nested rules on this scope
    nested = list(g.objects(scope, GEM.hasPolicyRule))
    render_rules(doc, g, nested, "Rules on this scope", level + 1)


def render_triples_appendix(doc, g, policy):
    """New page: the complete triple set for the policy and its nested subjects,
    as indented Turtle. Indentation reflects nesting: policy at depth 0; its
    rules, groups and coding rules at depth 1; rules on a group or coding rule
    at depth 2."""
    doc.add_page_break()
    doc.add_heading("RDF triples (Turtle)", level=1)
    add_muted(doc, "Complete triple set for the policy and the subjects nested "
                   "under it (rules, groups, coding rules; rules on a group or "
                   "coding rule nest one level deeper). Indentation reflects that "
                   "nesting. gem:workflowDescription triples are omitted.")
    doc.add_paragraph()

    lines, emitted = [], set()
    render_subject_ttl(g, policy, 0, lines, emitted)
    for r in sorted(g.objects(policy, GEM.hasPolicyRule), key=lambda x: sort_key(local(x))):
        render_subject_ttl(g, r, 1, lines, emitted)
    for grp in sorted(g.objects(policy, GEM.hasPolicyGroup), key=lambda x: sort_key(local(x))):
        render_subject_ttl(g, grp, 1, lines, emitted)
        for r in sorted(g.objects(grp, GEM.hasPolicyRule), key=lambda x: sort_key(local(x))):
            render_subject_ttl(g, r, 2, lines, emitted)
    for cr in sorted(g.objects(policy, GEM.hasPolicyCodingRule), key=lambda x: sort_key(local(x))):
        render_subject_ttl(g, cr, 1, lines, emitted)
        for r in sorted(g.objects(cr, GEM.hasPolicyRule), key=lambda x: sort_key(local(x))):
            render_subject_ttl(g, r, 2, lines, emitted)

    # drop the trailing blank line
    while lines and lines[-1] == "":
        lines.pop()
    for ln in lines:
        emit_code_line(doc, ln)


# --------------------------------------------------------------------------- #
# Main
# --------------------------------------------------------------------------- #

def load_graph(files_dir):
    g = rdflib.Graph()
    loaded = []
    for name in DEFAULT_FILES:
        path = os.path.join(files_dir, name)
        if os.path.exists(path):
            g.parse(path, format="turtle")
            loaded.append(name)
    if not loaded:
        # fall back: parse every .ttl in the directory
        for name in sorted(os.listdir(files_dir)):
            if name.endswith(".ttl"):
                g.parse(os.path.join(files_dir, name), format="turtle")
                loaded.append(name)
    return g, loaded


def resolve_policy(g, identifier):
    """Exact match on gem:identifier, whitespace-collapsed & case-insensitive."""
    want = re.sub(r"\s+", " ", identifier).strip().lower()
    hits = []
    for s, _, o in g.triples((None, GEM.identifier, None)):
        val = re.sub(r"\s+", " ", str(o)).strip().lower()
        if val == want:
            # must be a policy document, not a rule/CR/etc.
            types = set(local(t) for t in g.objects(s, RDF.type))
            if any(t.endswith("policy") or t in {"NCDpolicy", "LCDpolicy",
                    "ArticlePolicy", "TransmittalPolicy", "NCAdocument",
                    "CMSmanualChapter"} for t in types):
                hits.append(s)
    return hits


def build_doc(g, policy, loaded):
    doc = Document()
    set_letter(doc)
    style_base(doc)

    label = g.value(policy, GEM.prefLabel) or local(policy)
    ident = g.value(policy, GEM.identifier) or local(policy)
    ptype = ", ".join(sorted(local(t) for t in g.objects(policy, RDF.type)))

    # Title
    title = doc.add_heading(str(label), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph()
    r = sub.add_run(f"{ident}  ·  {ptype}")
    r.italic = True
    r.font.color.rgb = MUTED

    # --- Identity ---
    doc.add_heading("Policy", level=1)
    idrows = []
    for disp, pred in POLICY_IDENTITY_PREDS:
        vals = [obj_render(g, o) for o in g.objects(policy, pred)]
        if vals:
            idrows.append((disp, "; ".join(sorted(vals, key=sort_key))))
    src = g.value(policy, DC.source)
    if src:
        idrows.append(("Source", str(src)))
    add_kv_table(doc, idrows)

    desc = g.value(policy, GEM.description)
    if desc:
        doc.add_paragraph(str(desc))

    # --- Attributes & references (grouped object properties) ---
    buckets = grouped_objects(g, policy)
    if buckets:
        doc.add_heading("Attributes & references", level=2)
        for plabel in sorted(buckets):
            add_bullets(doc, buckets[plabel], label=plabel)

    # --- Policy-level rules ---
    rules = list(g.objects(policy, GEM.hasPolicyRule))
    if rules:
        render_rules(doc, g, rules, "Policy rules", 1)
    else:
        doc.add_heading("Policy rules", level=1)
        doc.add_paragraph("None.")

    # --- Policy Groups ---
    groups = list(g.objects(policy, GEM.hasPolicyGroup))
    doc.add_heading("Policy groups", level=1)
    if groups:
        for grp in sorted(groups, key=lambda x: sort_key(local(x))):
            render_scope(doc, g, grp, "Policy group", 2)
    else:
        doc.add_paragraph("None.")

    # --- Policy Coding Rules ---
    coding = list(g.objects(policy, GEM.hasPolicyCodingRule))
    doc.add_heading("Policy coding rules", level=1)
    if coding:
        for cr in sorted(coding, key=lambda x: sort_key(local(x))):
            render_scope(doc, g, cr, "Policy coding rule", 2)
    else:
        doc.add_paragraph("None.")

    # --- Footer note ---
    doc.add_paragraph()
    add_muted(doc, f"Generated from the GEM knowledge graph "
                   f"({', '.join(loaded)}). Curated readout: gem:workflowDescription omitted.")

    # --- Triples appendix (new page, after all other output) ---
    render_triples_appendix(doc, g, policy)
    return doc


def main():
    ap = argparse.ArgumentParser(description="Export a GEM policy to Word (.docx).")
    ap.add_argument("--identifier", required=True,
                    help='Policy identifier, e.g. "NCD 30.3".')
    ap.add_argument("--files-dir", default=".",
                    help="Directory holding the GEM .ttl files (default: cwd).")
    ap.add_argument("--output", default=None,
                    help="Output .docx path (default derived from identifier).")
    args = ap.parse_args()

    if not os.path.isdir(args.files_dir):
        sys.exit(f"ERROR: files-dir not found: {args.files_dir}")

    g, loaded = load_graph(args.files_dir)
    if not loaded:
        sys.exit(f"ERROR: no .ttl files found in {args.files_dir}")

    # Bind the namespace BEFORE any gem: query. Must precede resolve_policy:
    # under a wrong namespace it returns [] and the "no such policy" exit below
    # blames the identifier for what is really a namespace mismatch.
    gem_uri = bind_gem_namespace(g)
    if not any(g.triples((None, GEM.identifier, None))):
        sys.exit(f"ERROR: no gem:identifier triples under the detected namespace "
                 f"<{gem_uri}>. The graph in {args.files_dir} does not declare a "
                 f"usable `gem:` prefix; this is a namespace problem, not an "
                 f"identifier problem.")

    hits = resolve_policy(g, args.identifier)
    if not hits:
        sys.exit(f'ERROR: no policy individual with gem:identifier "{args.identifier}". '
                 f"Check the identifier (exact match only).")
    if len(hits) > 1:
        names = ", ".join(local(h) for h in hits)
        sys.exit(f'ERROR: identifier "{args.identifier}" matched multiple subjects: {names}')

    policy = hits[0]
    out = args.output
    if not out:
        safe = re.sub(r"[^A-Za-z0-9.]+", "_", args.identifier).strip("_")
        out = f"{safe}.docx"
    os.makedirs(os.path.dirname(os.path.abspath(out)), exist_ok=True)

    doc = build_doc(g, policy, loaded)
    doc.save(out)

    n_rules = len(list(g.objects(policy, GEM.hasPolicyRule)))
    n_groups = len(list(g.objects(policy, GEM.hasPolicyGroup)))
    n_coding = len(list(g.objects(policy, GEM.hasPolicyCodingRule)))
    print(f"Wrote {out}")
    print(f"  policy      : {local(policy)} ({g.value(policy, GEM.prefLabel)})")
    print(f"  policy rules: {n_rules}")
    print(f"  groups      : {n_groups}")
    print(f"  coding rules: {n_coding}")


if __name__ == "__main__":
    main()
